#!/usr/bin/env python3
"""
run_checks.py
Python-based rule checker for the Pleadings Checker project.

Creates a test Word document with content triggering all 34 rules,
runs simplified Python equivalents of each VBA rule, annotates
the document with comments, and saves the output as TEST1OUTPUT.

Usage:
    python3 run_checks.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
#  Issue dataclass
# ============================================================
class PleadingIssue:
    def __init__(self, rule_name, location, issue, suggestion, severity="error", auto_fix_safe=False):
        self.rule_name = rule_name
        self.location = location
        self.issue = issue
        self.suggestion = suggestion
        self.severity = severity
        self.auto_fix_safe = auto_fix_safe

    def __repr__(self):
        return f"[{self.severity}] {self.rule_name}: {self.issue} -> {self.suggestion} ({self.location})"


# ============================================================
#  Utility helpers
# ============================================================

def add_comment_to_paragraph(paragraph, comment_text):
    """Add a visible annotation as a bold red inline note at end of paragraph."""
    run = paragraph.add_run(f"  [{comment_text}]")
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(8)
    run.bold = True


def highlight_run_yellow(run):
    """Apply yellow highlight to a run."""
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def get_para_location(idx):
    return f"paragraph {idx + 1}"


# ============================================================
#  RULE IMPLEMENTATIONS (Python equivalents of VBA rules)
# ============================================================

# ── Rule 1: British Spelling ─────────────────────────────────
US_TO_UK = {
    "color": "colour", "colors": "colours", "colored": "coloured",
    "favor": "favour", "favors": "favours", "favorable": "favourable",
    "honor": "honour", "honors": "honours", "honorable": "honourable",
    "humor": "humour", "labor": "labour", "neighbor": "neighbour",
    "organize": "organise", "organized": "organised", "organizing": "organising",
    "organization": "organisation", "organizations": "organisations",
    "recognize": "recognise", "recognized": "recognised",
    "realize": "realise", "realized": "realised",
    "authorize": "authorise", "authorized": "authorised",
    "authorization": "authorisation",
    "minimize": "minimise", "maximise": "maximise",
    "analyze": "analyse", "analyzed": "analysed",
    "defense": "defence", "offense": "offence",
    # NB: "practice" (noun) is correct UK English; "practise" is the verb.
    # NB: "judgment" (no 'e') is standard in UK legal writing.
    # Neither is flagged here to avoid legal false positives.
    "center": "centre", "centers": "centres",
    "meter": "metre",
    "theater": "theatre",
    "catalog": "catalogue", "dialog": "dialogue",
    "acknowledgment": "acknowledgement",
    "aging": "ageing",
    "gray": "grey",
    "tire": "tyre",
    "skeptic": "sceptic",
    "maneuver": "manoeuvre",
    "pajamas": "pyjamas",
}


def rule01_british_spelling(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        words = re.findall(r'\b[a-zA-Z]+\b', text)
        for w in words:
            lw = w.lower()
            if lw in US_TO_UK:
                issues.append(PleadingIssue(
                    "british_spelling", get_para_location(idx),
                    f"US spelling detected: '{w}'",
                    f"Use '{US_TO_UK[lw]}' instead.",
                    "warning", True
                ))
    return issues


# ── Rule 2: Repeated Words ───────────────────────────────────
KNOWN_VALID_REPEATS = {"that", "had", "is", "was", "can"}


def rule02_repeated_words(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        words = re.findall(r'\b[a-zA-Z]+\b', para.text)
        for i in range(1, len(words)):
            if words[i].lower() == words[i - 1].lower():
                sev = "possible_error" if words[i].lower() in KNOWN_VALID_REPEATS else "error"
                issues.append(PleadingIssue(
                    "repeated_words", get_para_location(idx),
                    f"Repeated word: '{words[i]} {words[i]}'",
                    f"Remove duplicate '{words[i]}'.",
                    sev
                ))
    return issues


# ── Rule 3: Sequential Numbering ─────────────────────────────
def rule03_sequential_numbering(paragraphs):
    issues = []
    expected = 1
    for idx, para in enumerate(paragraphs):
        m = re.match(r'^(\d+)\.\s', para.text.strip())
        if m:
            num = int(m.group(1))
            if num != expected:
                issues.append(PleadingIssue(
                    "sequential_numbering", get_para_location(idx),
                    f"Expected clause {expected}, found {num}.",
                    f"Renumber to {expected}.",
                    "error"
                ))
            expected = num + 1
    return issues


# ── Rule 4: Heading Capitalisation ───────────────────────────
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def rule04_heading_capitalisation(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name in HEADING_STYLES:
            text = para.text.strip()
            if text and not text[0].isupper() and text[0].isalpha():
                issues.append(PleadingIssue(
                    "heading_capitalisation", get_para_location(idx),
                    f"Heading does not start with a capital letter.",
                    "Capitalise the first word of the heading.",
                    "error"
                ))
    return issues


# ── Rule 5: Custom Term Whitelist (no-op in Python) ──────────
def rule05_custom_term_whitelist(paragraphs):
    return []


# ── Rule 6: Paragraph Break Consistency ──────────────────────
def rule06_paragraph_break_consistency(paragraphs):
    issues = []
    spacings = []
    for para in paragraphs:
        if para.text.strip():
            fmt = para.paragraph_format
            after = fmt.space_after
            if after is not None:
                spacings.append(after)
    if spacings:
        common = max(set(spacings), key=spacings.count)
        for idx, para in enumerate(paragraphs):
            if para.text.strip():
                fmt = para.paragraph_format
                if fmt.space_after is not None and fmt.space_after != common:
                    issues.append(PleadingIssue(
                        "paragraph_break_consistency", get_para_location(idx),
                        "Inconsistent paragraph spacing.",
                        "Standardise paragraph spacing.",
                        "warning"
                    ))
    return issues


# ── Rule 7: Defined Terms ────────────────────────────────────
def rule07_defined_terms(paragraphs):
    issues = []
    full_text = "\n".join(p.text for p in paragraphs)
    # Find terms in quotes followed by "means"
    defined = re.findall(r'[\u201c"]([\w\s-]+?)[\u201d"]\s+means\b', full_text)
    # Find parenthetical definitions
    defined += re.findall(r'\((?:the\s+)?[\u201c"]([\w\s-]+?)[\u201d"]\)', full_text)
    for term in defined:
        term = term.strip()
        count = len(re.findall(re.escape(term), full_text))
        if count <= 1:
            issues.append(PleadingIssue(
                "defined_terms", "document level",
                f"Defined term '{term}' appears only at its definition.",
                "Remove unused defined term or add references.",
                "warning"
            ))
    return issues


# ── Rule 8: Clause Number Format ─────────────────────────────
def rule08_clause_number_format(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        if re.match(r'^[(\[]\d+[)\]]\s', para.text.strip()):
            issues.append(PleadingIssue(
                "clause_number_format", get_para_location(idx),
                "Clause number in brackets.",
                "Use 'N.' format without brackets.",
                "warning"
            ))
    return issues


# ── Rule 9: Date/Time Format ─────────────────────────────────
def rule09_date_time_format(paragraphs):
    issues = []
    us_date = re.compile(r'\b(\d{1,2})/(\d{1,2})/(\d{4})\b')
    for idx, para in enumerate(paragraphs):
        for m in us_date.finditer(para.text):
            month, day, year = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if month <= 12 and day <= 31:
                issues.append(PleadingIssue(
                    "date_time_format", get_para_location(idx),
                    f"Ambiguous or US-style date format: '{m.group()}'.",
                    "Use DD/MM/YYYY or written form (e.g. 1 January 2024).",
                    "warning"
                ))
    return issues


# ── Rule 10: Inline List Format ──────────────────────────────
def rule10_inline_list_format(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        # Detect inline lists with (a), (b), (c) pattern
        items = re.findall(r'\([a-z]\)\s[^;,]+[,]', text)
        if len(items) >= 2:
            issues.append(PleadingIssue(
                "inline_list_format", get_para_location(idx),
                "Inline list items separated by commas.",
                "Use semicolons between inline list items.",
                "warning"
            ))
    return issues


# ── Rule 11: Font Consistency ────────────────────────────────
def rule11_font_consistency(paragraphs):
    issues = []
    fonts = {}
    for para in paragraphs:
        for run in para.runs:
            fname = run.font.name
            if fname:
                fonts[fname] = fonts.get(fname, 0) + len(run.text)
    if len(fonts) > 1:
        dominant = max(fonts, key=fonts.get)
        for idx, para in enumerate(paragraphs):
            for run in para.runs:
                if run.font.name and run.font.name != dominant and len(run.text.strip()) > 0:
                    issues.append(PleadingIssue(
                        "font_consistency", get_para_location(idx),
                        f"Font '{run.font.name}' differs from dominant '{dominant}'.",
                        f"Use '{dominant}' throughout.",
                        "warning"
                    ))
                    break
    return issues


# ── Rule 12: Licence/License ─────────────────────────────────
def rule12_licence_license(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        # "license" as noun should be "licence" in UK English
        for m in re.finditer(r'\blicense\b', para.text, re.IGNORECASE):
            issues.append(PleadingIssue(
                "licence_license", get_para_location(idx),
                f"'license' detected (US spelling).",
                "Use 'licence' (noun) or 'license' (verb) per UK convention.",
                "warning"
            ))
    return issues


# ── Rule 13: Colour Formatting ───────────────────────────────
# NB: "color" is already caught by Rule 1 (british_spelling).
# This rule catches compound forms that Rule 1 misses (e.g. "colorful").
def rule13_colour_formatting(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        for m in re.finditer(r'\bcolor(?:ful|ing|ise|ize|ed|less)\b', para.text, re.IGNORECASE):
            issues.append(PleadingIssue(
                "colour_formatting", get_para_location(idx),
                f"US spelling '{m.group()}' detected.",
                f"Use '{m.group().replace('color', 'colour').replace('ize', 'ise')}'.",
                "warning", True
            ))
    return issues


# ── Rule 14: Slash Style ─────────────────────────────────────
def rule14_slash_style(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        # Detect "and/or" patterns
        for m in re.finditer(r'\b\w+/\w+\b', para.text):
            matched = m.group()
            if matched.lower() in ("and/or", "he/she", "his/her"):
                issues.append(PleadingIssue(
                    "slash_style", get_para_location(idx),
                    f"Slash construction '{matched}' detected.",
                    "Avoid slash constructions; use alternatives.",
                    "warning"
                ))
    return issues


# ── Rule 15: List Punctuation ────────────────────────────────
def rule15_list_punctuation(paragraphs):
    issues = []
    in_list = False
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if re.match(r'^[\u2022\-\*]\s', text) or re.match(r'^[a-z]\)\s', text):
            in_list = True
            if text and text[-1] not in '.;:,':
                issues.append(PleadingIssue(
                    "list_punctuation", get_para_location(idx),
                    "List item has no terminal punctuation.",
                    "Add semicolon or period to list item.",
                    "warning"
                ))
        else:
            in_list = False
    return issues


# ── Rule 16: Bracket Integrity ───────────────────────────────
BRACKET_PAIRS = {'(': ')', '[': ']', '{': '}'}
CLOSE_TO_OPEN = {v: k for k, v in BRACKET_PAIRS.items()}


def rule16_bracket_integrity(paragraphs):
    issues = []
    stack = []
    for idx, para in enumerate(paragraphs):
        for ch in para.text:
            if ch in BRACKET_PAIRS:
                stack.append((ch, idx))
            elif ch in CLOSE_TO_OPEN:
                if not stack:
                    issues.append(PleadingIssue(
                        "bracket_integrity", get_para_location(idx),
                        f"Unmatched closing bracket '{ch}'.",
                        "Add or correct matching opening bracket.",
                        "error"
                    ))
                else:
                    open_ch, open_idx = stack.pop()
                    if BRACKET_PAIRS.get(open_ch) != ch:
                        issues.append(PleadingIssue(
                            "bracket_integrity", get_para_location(idx),
                            f"Mismatched brackets: '{open_ch}' closed with '{ch}'.",
                            "Correct bracket pairing.",
                            "error"
                        ))
    for open_ch, open_idx in stack:
        issues.append(PleadingIssue(
            "bracket_integrity", get_para_location(open_idx),
            f"Unmatched opening bracket '{open_ch}'.",
            "Add corresponding closing bracket.",
            "error"
        ))
    return issues


# ── Rule 17: Quotation Mark Consistency ──────────────────────
def rule17_quotation_mark_consistency(paragraphs):
    issues = []
    double_smart = 0
    double_straight = 0
    for para in paragraphs:
        for ch in para.text:
            if ch in ('\u201c', '\u201d'):
                double_smart += 1
            elif ch == '"':
                double_straight += 1
    if double_smart > 0 and double_straight > 0:
        issues.append(PleadingIssue(
            "quotation_mark_consistency", "document level",
            f"Mixed quotation marks: {double_smart} smart, {double_straight} straight.",
            "Standardise quotation marks.",
            "warning"
        ))
    return issues


# ── Rule 18: Page Range (no-op - config only) ────────────────
def rule18_page_range(paragraphs):
    return []


# ── Rule 19: Currency/Number Format ──────────────────────────
def rule19_currency_number_format(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        # Detect $
        for m in re.finditer(r'\$[\d,]+\.?\d*', para.text):
            issues.append(PleadingIssue(
                "currency_number_format", get_para_location(idx),
                f"Dollar sign '{m.group()}' detected.",
                "Use GBP (£) or specify currency explicitly.",
                "warning"
            ))
    return issues


# ── Rule 20: Footnote Integrity ──────────────────────────────
def rule20_footnote_integrity(paragraphs):
    issues = []
    full_text = "\n".join(p.text for p in paragraphs)
    # Check for footnote markers like [1] [2] etc and verify sequential
    # Filter out year-like numbers (4+ digits) and case citations
    markers = [m for m in re.findall(r'\[(\d+)\]', full_text) if len(m) <= 2]
    expected = 1
    for m in markers:
        num = int(m)
        if num != expected:
            issues.append(PleadingIssue(
                "footnote_integrity", "document level",
                f"Footnote reference [{num}] out of sequence (expected [{expected}]).",
                f"Renumber footnote to [{expected}].",
                "error"
            ))
        expected = num + 1
    return issues


# ── Rule 21: Title Formatting ────────────────────────────────
def rule21_title_formatting(paragraphs):
    issues = []
    if paragraphs:
        style_name = paragraphs[0].style.name if paragraphs[0].style else ""
        if "Title" not in style_name and "Heading" not in style_name:
            text = paragraphs[0].text.strip()
            if text and len(text) < 200:
                issues.append(PleadingIssue(
                    "title_formatting", "paragraph 1",
                    "First paragraph does not use Title/Heading style.",
                    "Apply 'Title' or 'Heading 1' style.",
                    "warning"
                ))
    return issues


# ── Rule 22: Brand Name Enforcement ──────────────────────────
BRAND_RULES = {
    "PwC": ["pwc", "Pwc", "PWC", "PricewaterhouseCoopers", "pricewaterhousecoopers"],
    "Deloitte": ["deloitte", "DELOITTE"],
    "HMRC": ["hmrc", "Hmrc", "H.M.R.C.", "HM Revenue"],
    "FCA": ["fca", "Fca", "F.C.A."],
    "EY": ["ey", "Ey", "Ernst & Young", "ernst & young", "E&Y", "E.Y."],
    "KPMG": ["kpmg", "Kpmg", "K.P.M.G."],
}


def rule22_brand_name_enforcement(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        for correct, variants in BRAND_RULES.items():
            for variant in variants:
                pattern = re.compile(r'\b' + re.escape(variant) + r'\b', re.IGNORECASE if variant.islower() else 0)
                for m in pattern.finditer(text):
                    if m.group() != correct:
                        issues.append(PleadingIssue(
                            "brand_name_enforcement", get_para_location(idx),
                            f"Brand name variant '{m.group()}' detected.",
                            f"Use '{correct}'.",
                            "warning", True
                        ))
    return issues


# ── Rule 23: Phrase Consistency ──────────────────────────────
def rule23_phrase_consistency(paragraphs):
    issues = []
    full_text = "\n".join(p.text for p in paragraphs)
    phrase_groups = [
        (["notwithstanding", "despite", "in spite of"], "despite"),
        (["hereinafter", "hereafter"], "hereafter"),
        (["pursuant to", "in accordance with"], "in accordance with"),
    ]
    for group, preferred in phrase_groups:
        found = {}
        for phrase in group:
            cnt = len(re.findall(re.escape(phrase), full_text, re.IGNORECASE))
            if cnt > 0:
                found[phrase] = cnt
        if len(found) > 1:
            issues.append(PleadingIssue(
                "phrase_consistency", "document level",
                f"Inconsistent phrasing: {', '.join(found.keys())} all used.",
                f"Standardise to '{preferred}' or choose one form.",
                "warning"
            ))
    return issues


# ── Rule 24: Footnotes Not Endnotes ──────────────────────────
def rule24_footnotes_not_endnotes(paragraphs):
    issues = []
    full_text = "\n".join(p.text for p in paragraphs)
    if re.search(r'\bendnote\b', full_text, re.IGNORECASE):
        issues.append(PleadingIssue(
            "footnotes_not_endnotes", "document level",
            "Reference to endnotes detected.",
            "Use footnotes rather than endnotes.",
            "error"
        ))
    return issues


# ── Rule 25: Footnote Terminal Full Stop ─────────────────────
def rule25_footnote_terminal_full_stop(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if text.startswith("[fn]"):
            fn_text = text[4:].strip()
            if fn_text and not fn_text.endswith('.'):
                issues.append(PleadingIssue(
                    "footnote_terminal_full_stop", get_para_location(idx),
                    "Footnote does not end with a full stop.",
                    "Add a period at the end of the footnote.",
                    "warning"
                ))
    return issues


# ── Rule 26: Footnote Initial Capital ────────────────────────
ALLOWED_LOWERCASE_STARTS = {"ibid", "eg", "ie", "cf", "supra", "infra"}


def rule26_footnote_initial_capital(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if text.startswith("[fn]"):
            fn_text = text[4:].strip()
            if fn_text:
                # Extract first token (may include dots, e.g. "e.g.")
                first_token = re.match(r'([a-zA-Z]+(?:\.[a-zA-Z])*\.?)', fn_text)
                if first_token:
                    token = first_token.group(1)
                    # Normalise: strip dots for comparison
                    normalised = token.replace('.', '').lower()
                    if token[0].islower() and normalised not in ALLOWED_LOWERCASE_STARTS:
                        issues.append(PleadingIssue(
                            "footnote_initial_capital", get_para_location(idx),
                            f"Footnote starts with lowercase '{token}'.",
                            "Capitalise the first word of the footnote.",
                            "warning"
                        ))
    return issues


# ── Rule 27: Footnote Abbreviation Dictionary ────────────────
APPROVED_ABBREVS = {
    "art", "arts", "ch", "chs", "c", "cc", "cl", "cls", "cp", "cf",
    "ed", "eds", "edn", "edns", "eg", "etc", "f", "ff", "fn", "fns",
    "ibid", "ie", "ms", "mss", "n", "nn", "no", "p", "pp", "para",
    "paras", "pt", "reg", "regs", "r", "rr", "sch", "s", "ss",
    "trans", "vol", "vols"
}
UNAPPROVED_ABBREVS = {"pgs": "pp", "sec": "s", "secs": "ss", "sect": "s", "sects": "ss"}


def rule27_footnote_abbreviation_dictionary(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if text.startswith("[fn]"):
            fn_text = text[4:].strip()
            tokens = fn_text.split()
            for token in tokens:
                clean = token.strip('.,;:()[]')
                lc = clean.lower()
                # Check unapproved variants
                if lc in UNAPPROVED_ABBREVS:
                    issues.append(PleadingIssue(
                        "footnote_abbreviation_dictionary", get_para_location(idx),
                        f"Unapproved abbreviation '{clean}'.",
                        f"Use '{UNAPPROVED_ABBREVS[lc]}' instead.",
                        "warning"
                    ))
                # Check dotted forms of approved abbreviations
                elif '.' in token:
                    stripped = clean.rstrip('.')
                    no_dots = stripped.replace('.', '')
                    if no_dots.lower() in APPROVED_ABBREVS and len(no_dots) > 0:
                        issues.append(PleadingIssue(
                            "footnote_abbreviation_dictionary", get_para_location(idx),
                            f"Dotted abbreviation '{token.strip(',;:()[]')}' detected.",
                            f"Use '{no_dots}' without dots.",
                            "warning"
                        ))
    return issues


# ── Rule 28: Mandated Legal Term Forms ───────────────────────
MANDATED_TERMS = {
    "Solicitor-General": "Solicitor General",
    "Attorney-General": "Attorney General",
}


def rule28_mandated_legal_term_forms(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        for correct, wrong in MANDATED_TERMS.items():
            for m in re.finditer(re.escape(wrong), text, re.IGNORECASE):
                issues.append(PleadingIssue(
                    "mandated_legal_term_forms", get_para_location(idx),
                    f"Unhyphenated form '{m.group()}' detected.",
                    f"Use '{correct}'.",
                    "warning", False
                ))
    return issues


# ── Rule 29: Always Capitalise Terms ─────────────────────────
ALWAYS_CAPITALISE = {
    "Prime Minister": "prime minister",
    "Law Lords": "law lords",
    "Lord Chancellor": "lord chancellor",
    "Master of the Rolls": "master of the rolls",
    "Lord Chief Justice": "lord chief justice",
}


def rule29_always_capitalise_terms(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        for correct, wrong in ALWAYS_CAPITALISE.items():
            for m in re.finditer(re.escape(wrong), text):
                issues.append(PleadingIssue(
                    "always_capitalise_terms", get_para_location(idx),
                    f"'{m.group()}' should be capitalised.",
                    f"Use '{correct}'.",
                    "warning", False
                ))
    return issues


# ── Rule 30: Anglicised Terms Not Italic ─────────────────────
ANGLICISED_TERMS = [
    "prima facie", "per se", "de facto", "de jure", "ex parte",
    "bona fide", "inter alia", "ultra vires", "vice versa",
    "a priori", "a fortiori", "stare decisis", "obiter dicta",
    "ratio decidendi", "mutatis mutandis", "quid pro quo",
]


def rule30_anglicised_terms_not_italic(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        for run in para.runs:
            if run.italic:
                text = run.text
                for term in ANGLICISED_TERMS:
                    if term.lower() in text.lower():
                        issues.append(PleadingIssue(
                            "anglicised_terms_not_italic", get_para_location(idx),
                            f"Anglicised term '{term}' is in italics.",
                            f"Set '{term}' in roman type.",
                            "warning"
                        ))
    return issues


# ── Rule 31: Foreign Names Not Italic ────────────────────────
FOREIGN_NAMES = [
    "Cour de cassation", "Bundesgerichtshof", "Conseil d'Etat",
    "Tribunal de grande instance", "Corte di Cassazione",
]


def rule31_foreign_names_not_italic(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        for run in para.runs:
            if run.italic:
                text = run.text
                for name in FOREIGN_NAMES:
                    if name.lower() in text.lower():
                        issues.append(PleadingIssue(
                            "foreign_names_not_italic", get_para_location(idx),
                            f"Foreign institution name '{name}' is in italics.",
                            f"Set '{name}' in roman type.",
                            "warning"
                        ))
    return issues


# ── Rule 32: Single Quotes Default ───────────────────────────
def rule32_single_quotes_default(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        style_name = para.style.name if para.style else ""
        if "Quote" in style_name:
            continue
        if '\u201c' in text or '\u201d' in text:
            issues.append(PleadingIssue(
                "single_quotes_default", get_para_location(idx),
                "Outer quotation uses double quotes.",
                "Use single quotes (\u2018...\u2019) for outer quotations.",
                "warning"
            ))
    return issues


# ── Rule 33: Smart Quote Consistency ─────────────────────────
def rule33_smart_quote_consistency(paragraphs):
    issues = []
    straight_count = 0
    curly_count = 0
    full_text = "\n".join(p.text for p in paragraphs)
    for ch in full_text:
        code = ord(ch)
        if code == 34:  # straight double
            straight_count += 1
        elif code in (8220, 8221):  # curly double
            curly_count += 1
        elif code == 39:  # straight single - check if apostrophe
            straight_count += 1
        elif code in (8216, 8217):  # curly single
            curly_count += 1
    if straight_count > 0 and curly_count > 0:
        issues.append(PleadingIssue(
            "smart_quote_consistency", "document level",
            f"Mixed quote styles: {straight_count} straight, {curly_count} curly.",
            "Use curly quotation marks consistently.",
            "warning"
        ))
    return issues


# ── Rule 34: Spell Out Under Ten ─────────────────────────────
STRUCTURAL_PREFIXES = {"para", "paragraph", "section", "clause", "article",
                       "page", "rule", "part", "schedule", "chapter", "no",
                       "item", "annex", "table", "figure", "reg"}


def rule34_spell_out_under_ten(paragraphs):
    issues = []
    for idx, para in enumerate(paragraphs):
        text = para.text
        # Skip numbered list items (e.g. "1. The first clause")
        if re.match(r'^\d+\.\s', text.strip()):
            continue
        # Skip bracket-format list items
        if re.match(r'^[\(\[]\d+[\)\]]\s', text.strip()):
            continue
        # Skip footnote-prefix paragraphs
        if text.strip().startswith("[fn]"):
            continue
        for m in re.finditer(r'(?<!\d)\b([1-9])\b(?!\d)', text):
            digit = m.group(1)
            start = m.start()
            # Check preceding word for structural reference
            before = text[:start].rstrip()
            prev_word = before.split()[-1].lower().rstrip('.') if before.split() else ""
            if prev_word in STRUCTURAL_PREFIXES:
                continue
            # Check if part of a range (digit-digit or digit\u2013digit)
            after = text[m.end():m.end() + 2]
            if after and after[0] in ('-', '\u2013'):
                continue
            before_ch = text[start - 1] if start > 0 else ''
            if before_ch in ('-', '\u2013'):
                continue
            # Skip numbers inside square brackets (footnote refs like [1])
            if start > 0 and text[start - 1] == '[':
                continue
            issues.append(PleadingIssue(
                "spell_out_under_ten", get_para_location(idx),
                f"Number '{digit}' should be spelled out.",
                f"Write '{['zero','one','two','three','four','five','six','seven','eight','nine'][int(digit)]}' instead.",
                "warning"
            ))
    return issues


# ============================================================
#  TEST DOCUMENT CREATION
# ============================================================

def create_test_document(path):
    """Create a test Word document with content that triggers all 34 rules."""
    doc = Document()

    # Title (deliberately not using Title style -> Rule 21)
    p = doc.add_paragraph("sample pleading document for review")
    p.style = doc.styles['Normal']

    # Rule 1: US spellings
    doc.add_paragraph(
        "The color of the organization was favorable to the defense. "
        "We must analyze this matter and recognize the gray area."
    )

    # Rule 2: Repeated words
    doc.add_paragraph("The the court held that the decision was was correct and proper.")

    # Rule 3: Sequential numbering (deliberate gap: 1, 2, 4)
    doc.add_paragraph("1. The first clause of the agreement.")
    doc.add_paragraph("2. The second clause concerning obligations.")
    doc.add_paragraph("4. The fourth clause is out of sequence.")

    # Rule 4: Heading capitalisation (lowercase heading)
    h = doc.add_heading("background and context", level=2)

    # Rule 6: Inconsistent paragraph spacing (varies throughout)
    p = doc.add_paragraph("This paragraph has standard spacing.")
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph("This paragraph has different spacing.")
    p.paragraph_format.space_after = Pt(6)

    # Rule 7: Defined terms
    doc.add_paragraph(
        '\u201cSpecial Purpose Vehicle\u201d means a company established for a specific transaction.'
    )

    # Rule 8: Clause number in brackets
    doc.add_paragraph("(1) The first sub-clause uses bracket format.")

    # Rule 9: Date format issues
    doc.add_paragraph("The agreement was signed on 03/15/2024 and amended on 12/01/2023.")

    # Rule 10: Inline list with commas instead of semicolons
    doc.add_paragraph(
        "The parties must: (a) file the documents, (b) serve notice, (c) attend the hearing, and (d) comply with orders."
    )

    # Rule 11: Font inconsistency
    p = doc.add_paragraph()
    run1 = p.add_run("This text is in Times New Roman. ")
    run1.font.name = "Times New Roman"
    run2 = p.add_run("But this text is in Arial.")
    run2.font.name = "Arial"

    # Rule 12: License vs Licence
    doc.add_paragraph("The license was granted under the terms of the agreement.")

    # Rule 13: Color formatting (covered by Rule 1 too, but explicit)
    doc.add_paragraph("The document uses color coding for different sections.")

    # Rule 14: Slash style
    doc.add_paragraph("The claimant and/or respondent must file his/her response.")

    # Rule 15: List punctuation (missing terminal punctuation)
    doc.add_paragraph("The following items are required:")
    doc.add_paragraph("- First item without punctuation")
    doc.add_paragraph("- Second item also missing punctuation")
    doc.add_paragraph("- Third item has a period.")

    # Rule 16: Bracket integrity (unmatched bracket)
    doc.add_paragraph("The court held (in its reasoning that the claim [was valid.")

    # Rule 17 & 33: Quotation mark consistency (mixed styles)
    doc.add_paragraph(
        'The witness stated \u201cI saw the incident\u201d and later said "I was certain".'
    )

    # Rule 19: Currency format
    doc.add_paragraph("The damages were assessed at $50,000 by the tribunal.")

    # Rule 20: Footnote integrity (out of sequence)
    doc.add_paragraph("The court noted [1] the precedent and [3] the statute.")

    # Rule 22: Brand name enforcement
    doc.add_paragraph("The audit was conducted by pwc and reviewed by hmrc officials.")

    # Rule 23: Phrase consistency (notwithstanding vs despite; pursuant to vs in accordance with)
    doc.add_paragraph(
        "Notwithstanding the terms above, and despite earlier agreements, "
        "the parties acted in accordance with the new protocol."
    )
    doc.add_paragraph(
        "Pursuant to clause 5, the obligations remain in full force."
    )

    # Rule 24: Footnotes not endnotes
    doc.add_paragraph("See endnote 3 for further discussion of this point.")

    # Rule 25 & 26 & 27: Footnote issues (simulated with [fn] prefix)
    doc.add_paragraph("[fn] See Smith v Jones [2020] UKSC 1.")  # correct
    doc.add_paragraph("[fn] see https://example.com")  # no stop, lowercase
    doc.add_paragraph("[fn] e.g. discussion in paras. 4-6 of the pgs 12-13.")  # dotted abbrevs

    # Rule 28: Mandated legal term forms
    doc.add_paragraph("The Solicitor General and Attorney General appeared for the Crown.")

    # Rule 29: Always capitalise terms
    doc.add_paragraph("The prime minister addressed the law lords about the matter.")

    # Rule 30: Anglicised terms in italics
    p = doc.add_paragraph()
    run_normal = p.add_run("The evidence was ")
    run_italic = p.add_run("prima facie")
    run_italic.italic = True
    run_after = p.add_run(" sufficient to establish the claim ")
    run_italic2 = p.add_run("per se")
    run_italic2.italic = True
    run_after2 = p.add_run(".")

    # Rule 31: Foreign names in italics
    p = doc.add_paragraph()
    run_normal = p.add_run("The decision of the ")
    run_italic = p.add_run("Cour de cassation")
    run_italic.italic = True
    run_after = p.add_run(" was cited with approval.")

    # Rule 32: Double quotes as outer (using smart double quotes)
    doc.add_paragraph(
        "He argued \u201cthe principle is well established\u201d before the court."
    )

    # Rule 34: Spell out numbers under 10
    doc.add_paragraph("There were 7 issues raised and 3 witnesses called at the hearing.")

    doc.save(path)
    print(f"Test document created: {path}")
    return path


# ============================================================
#  RUN ALL CHECKS
# ============================================================

ALL_RULES = [
    ("Rule 01", rule01_british_spelling),
    ("Rule 02", rule02_repeated_words),
    ("Rule 03", rule03_sequential_numbering),
    ("Rule 04", rule04_heading_capitalisation),
    ("Rule 05", rule05_custom_term_whitelist),
    ("Rule 06", rule06_paragraph_break_consistency),
    ("Rule 07", rule07_defined_terms),
    ("Rule 08", rule08_clause_number_format),
    ("Rule 09", rule09_date_time_format),
    ("Rule 10", rule10_inline_list_format),
    ("Rule 11", rule11_font_consistency),
    ("Rule 12", rule12_licence_license),
    ("Rule 13", rule13_colour_formatting),
    ("Rule 14", rule14_slash_style),
    ("Rule 15", rule15_list_punctuation),
    ("Rule 16", rule16_bracket_integrity),
    ("Rule 17", rule17_quotation_mark_consistency),
    ("Rule 18", rule18_page_range),
    ("Rule 19", rule19_currency_number_format),
    ("Rule 20", rule20_footnote_integrity),
    ("Rule 21", rule21_title_formatting),
    ("Rule 22", rule22_brand_name_enforcement),
    ("Rule 23", rule23_phrase_consistency),
    ("Rule 24", rule24_footnotes_not_endnotes),
    ("Rule 25", rule25_footnote_terminal_full_stop),
    ("Rule 26", rule26_footnote_initial_capital),
    ("Rule 27", rule27_footnote_abbreviation_dictionary),
    ("Rule 28", rule28_mandated_legal_term_forms),
    ("Rule 29", rule29_always_capitalise_terms),
    ("Rule 30", rule30_anglicised_terms_not_italic),
    ("Rule 31", rule31_foreign_names_not_italic),
    ("Rule 32", rule32_single_quotes_default),
    ("Rule 33", rule33_smart_quote_consistency),
    ("Rule 34", rule34_spell_out_under_ten),
]


def run_all_checks(input_path, output_path):
    """Run all 34 rules on the input document and save annotated output."""
    doc = Document(input_path)
    paragraphs = list(doc.paragraphs)

    all_issues = []
    rule_counts = {}

    print("=" * 60)
    print("  Pleadings Checker - All Rules")
    print("=" * 60)

    for rule_name, rule_fn in ALL_RULES:
        issues = rule_fn(paragraphs)
        all_issues.extend(issues)
        rule_counts[rule_name] = len(issues)
        if issues:
            print(f"  {rule_name}: {len(issues)} issue(s)")
            for issue in issues:
                print(f"    [{issue.severity}] {issue.issue}")
        else:
            print(f"  {rule_name}: OK")

    print("=" * 60)
    total = len(all_issues)
    print(f"  TOTAL ISSUES: {total}")
    print("=" * 60)

    # Annotate the document: add comments to paragraphs with issues
    # Group issues by paragraph location
    para_issues = {}
    for issue in all_issues:
        loc = issue.location
        if loc not in para_issues:
            para_issues[loc] = []
        para_issues[loc].append(issue)

    # Add annotations
    for loc, issues in para_issues.items():
        m = re.match(r'paragraph (\d+)', loc)
        if m:
            para_idx = int(m.group(1)) - 1
            if 0 <= para_idx < len(paragraphs):
                # Highlight existing runs yellow
                for run in paragraphs[para_idx].runs:
                    highlight_run_yellow(run)
                # Add comment summary
                comment_parts = []
                for iss in issues:
                    comment_parts.append(f"{iss.rule_name}: {iss.suggestion}")
                comment_text = " | ".join(comment_parts[:3])
                if len(comment_parts) > 3:
                    comment_text += f" | +{len(comment_parts) - 3} more"
                add_comment_to_paragraph(paragraphs[para_idx], comment_text)

    # Also handle document-level issues - add a summary paragraph
    doc_level = [i for i in all_issues if i.location == "document level"]
    if doc_level:
        doc.add_paragraph("")  # spacer
        p = doc.add_paragraph()
        run = p.add_run("=== DOCUMENT-LEVEL ISSUES ===")
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        for iss in doc_level:
            p = doc.add_paragraph()
            run = p.add_run(f"[{iss.rule_name}] {iss.issue} -> {iss.suggestion}")
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.size = Pt(9)

    doc.save(output_path)
    print(f"\nAnnotated output saved: {output_path}")
    return all_issues


# ============================================================
#  MAIN
# ============================================================

if __name__ == "__main__":
    base_dir = Path(__file__).parent
    test_doc = base_dir / "test_pleading.docx"
    output_doc = base_dir / "test_pleading_TEST1OUTPUT.docx"

    # Step 1: Create test document
    create_test_document(str(test_doc))

    # Step 2: Run all checks and produce annotated output
    issues = run_all_checks(str(test_doc), str(output_doc))

    print(f"\nDone. {len(issues)} total issues found across all rules.")
